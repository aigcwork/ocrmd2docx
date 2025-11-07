#!/usr/bin/env python3
"""
创建一个基本的reference.docx文件，作为Pandoc的中文模板
"""

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    
    # 创建新文档
    doc = Document()
    
    # 设置默认字体为支持中文的字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'  # 设置中文字体
    font.size = Pt(10.5)
    
    # 设置文档标题
    title = doc.add_heading('Markdown转Word模板', 0)
    
    # 添加一些示例文本
    doc.add_paragraph('这是一个用于Pandoc转换的Word模板文档。')
    doc.add_paragraph('此模板确保生成的Word文档能正确显示中文内容。')
    
    # 添加一个示例公式
    doc.add_paragraph('示例公式: E = mc²')
    
    # 保存文档
    doc.save('reference.docx')
    print("成功创建reference.docx文件")
    
except ImportError:
    print("需要安装python-docx库: pip install python-docx")
    # 创建一个简单的空文档作为占位符
    with open('reference.docx', 'wb') as f:
        # 写入一个最小的DOCX文件头
        f.write(b'PK\x03\x04\x14\x00\x06\x00\x08\x00\x00\x00!\x00')
    print("创建了基本的reference.docx占位符文件")